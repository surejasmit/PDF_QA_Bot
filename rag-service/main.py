from fastapi import FastAPI, HTTPException, Request
from pydantic import BaseModel, Field, validator
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_core.documents import Document
from langchain_core.prompts import PromptTemplate
from groq import Groq
from dotenv import load_dotenv
from transformers import (
    AutoConfig,
    AutoTokenizer,
    AutoModelForSeq2SeqLM,
    AutoModelForCausalLM,
)
from slowapi import Limiter
from slowapi.util import get_remote_address
from pathlib import Path
from contextlib import asynccontextmanager
import uvicorn
import torch
import os
import re
import time
import docx
import asyncio

# ===============================
# APP SETUP
# ===============================
load_dotenv()

BASE_DIR = Path(__file__).resolve().parent.parent
UPLOAD_DIR = (BASE_DIR / "uploads").resolve()

cleanup_task = None

@asynccontextmanager
async def lifespan(app: FastAPI):
    global cleanup_task
    cleanup_task = asyncio.create_task(background_cleanup())
    yield
    cleanup_task.cancel()
    try:
        await cleanup_task
    except asyncio.CancelledError:
        pass

app = FastAPI(lifespan=lifespan)
limiter = Limiter(key_func=get_remote_address)
app.state.limiter = limiter

# ===============================
# CONFIG
# ===============================
HF_GENERATION_MODEL = os.getenv("HF_GENERATION_MODEL", "google/flan-t5-small")
LLM_GENERATION_TIMEOUT = int(os.getenv("LLM_GENERATION_TIMEOUT", "30"))
SESSION_TIMEOUT = 3600

sessions = {}

embedding_model = HuggingFaceEmbeddings(
    model_name="sentence-transformers/all-MiniLM-L6-v2"
)

SUPPORTED_EXTENSIONS = [".pdf", ".docx", ".txt", ".md"]

generation_tokenizer = None
generation_model = None
generation_is_encoder_decoder = False

# ===============================
# TEXT CLEANING
# ===============================
def normalize_spaced_text(text: str) -> str:
    pattern = r"\b(?:[A-Za-z] ){2,}[A-Za-z]\b"
    return re.sub(pattern, lambda m: m.group(0).replace(" ", ""), text)


def normalize_answer(text: str) -> str:
    """
    Post-processes the LLM-generated answer.
    """
    text = normalize_spaced_text(text)
    text = re.sub(r"^(Answer[^:]*:|Context:|Question:)\s*", "", text, flags=re.I)
    return text.strip()


# ===============================
# DOCUMENT LOADERS
# ===============================
def load_pdf(file_path: str):
    return PyPDFLoader(file_path).load()


def load_txt(file_path: str):
    with open(file_path, "r", encoding="utf-8") as f:
        return [Document(page_content=f.read())]


def load_docx(file_path: str):
    doc = docx.Document(file_path)
    text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    return [Document(page_content=text)]


def load_document(file_path: str):
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return load_pdf(file_path)
    elif ext == ".docx":
        return load_docx(file_path)
    elif ext in [".txt", ".md"]:
        return load_txt(file_path)
    else:
        raise ValueError("Unsupported file format")


# ===============================
# MODEL LOADING
# ===============================
def load_generation_model():
    global generation_model, generation_tokenizer, generation_is_encoder_decoder

    if generation_model:
        return generation_tokenizer, generation_model, generation_is_encoder_decoder

    config = AutoConfig.from_pretrained(HF_GENERATION_MODEL)
    generation_is_encoder_decoder = bool(config.is_encoder_decoder)

    generation_tokenizer = AutoTokenizer.from_pretrained(HF_GENERATION_MODEL)

    if generation_is_encoder_decoder:
        generation_model = AutoModelForSeq2SeqLM.from_pretrained(HF_GENERATION_MODEL)
    else:
        generation_model = AutoModelForCausalLM.from_pretrained(HF_GENERATION_MODEL)

    if torch.cuda.is_available():
        generation_model = generation_model.to("cuda")

    generation_model.eval()
    return generation_tokenizer, generation_model, generation_is_encoder_decoder


def generate_response(prompt: str, max_new_tokens: int):
    tokenizer, model, is_enc = load_generation_model()
    device = next(model.parameters()).device

    encoded = tokenizer(prompt, return_tensors="pt", truncation=True, max_length=2048)
    encoded = {k: v.to(device) for k, v in encoded.items()}

    output = model.generate(
        **encoded,
        max_new_tokens=max_new_tokens,
        do_sample=False,
        pad_token_id=tokenizer.pad_token_id or tokenizer.eos_token_id,
    )

    if is_enc:
        return tokenizer.decode(output[0], skip_special_tokens=True)

    return tokenizer.decode(
        output[0][encoded["input_ids"].shape[1]:],
        skip_special_tokens=True,
    )


# ===============================
# REQUEST MODELS
# ===============================
class DocumentPath(BaseModel):
    filePath: str
    session_id: str


class AskRequest(BaseModel):
    question: str = Field(..., min_length=1)
    session_id: str
    history: list = []

    @validator("question")
    def validate_question(cls, v):
        if not v.strip():
            raise ValueError("Empty question")
        return v.strip()


class SummarizeRequest(BaseModel):
    session_id: str
    pdf: str | None = None



# ===============================
# SESSION CLEANUP
# ===============================
def cleanup_sessions():
    now = time.time()
    expired = [k for k, v in sessions.items()
               if now - v["last"] > SESSION_TIMEOUT]
    for k in expired:
        del sessions[k]

async def background_cleanup():
    while True:
        await asyncio.sleep(300)
        cleanup_sessions()


# ===============================
# PROCESS DOCUMENT
# ===============================
@app.post("/process")
@limiter.limit("15/15 minutes")
def process_doc(request: Request, data: DocumentPath):
    cleanup_sessions()

    if not os.path.exists(data.filePath):
        raise HTTPException(404, "File not found")

    docs = load_document(data.filePath)

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=150,
    )
    chunks = splitter.split_documents(docs)

    vectorstore = FAISS.from_documents(chunks, embedding_model)

    sessions[data.session_id] = {
        "vectorstore": vectorstore,
        "last": time.time(),
    }

    return {"message": "Processed successfully"}


# ===============================
# ASK
# ===============================
@app.post("/ask")
@limiter.limit("60/15 minutes")
def ask(request: Request, data: AskRequest):
    cleanup_sessions()

    session = sessions.get(data.session_id)
    if not session:
        return {"answer": "Session expired", "confidence_score": 0}

    vectorstore = session["vectorstore"]

    docs = vectorstore.similarity_search_with_score(data.question, k=4)

    if not docs:
        return {"answer": "No relevant info", "confidence_score": 0}

    context = "\n\n".join(d.page_content for d, _ in docs)

    prompt = f"""
    Answer ONLY using the context.

    Context:
    {context}

    Question:
    {data.question}

    user_prompt = CCC_PROMPT.format(
        context=context,
        question=question_with_history,
    )

    answer = generate_response(prompt, 150)

    session["last"] = time.time()

    return {"answer": normalize_answer(answer), "confidence_score": 85}


# ===============================
# SUMMARIZE
# ===============================
@app.post("/summarize")
def summarize(data: SummarizeRequest):
    session = sessions.get(data.session_id)
    if not session:
        return {"summary": "Session expired"}

    docs = session["vectorstore"].similarity_search("summary", k=6)

    context = "\n".join(d.page_content for d in docs)

    prompt = f"Summarize in bullet points:\n{context}"

    summary = generate_response(prompt, 220)

    return {"summary": normalize_answer(summary)}


# ===============================
# START
# ===============================
if __name__ == "__main__":
    uvicorn.run("main:app", host="0.0.0.0", port=5000, reload=True)